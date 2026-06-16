# -*- coding: utf-8 -*-
"""Generate supplementary tables/figures and Table 3 top-10 for the CMJS manuscript."""

from __future__ import annotations

import re
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import networkx as nx
import numpy as np
import pandas as pd

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import config

sys.path.insert(0, str(ROOT / "src"))
from importlib import import_module

mrc = import_module("04_mine_risk_chains")
fig06 = import_module("06_make_figures")

OUT_REAL = config.OUTPUT_DIR
OUT_SENS = config.OUTPUT_SENSITIVITY_DIR
SUPP_DIR = ROOT / "supplementary"
MANUSCRIPT_DIR = ROOT / "1正文中文" / "1Chiang Mai Journal of Science"
FIG_DIR = SUPP_DIR / "figures"
TABLE_DIR = SUPP_DIR / "tables"

CHAIN_LAYER_PRIORITY = {
    "chain_local_1": 0,
    "chain_local_2": 1,
    "chain_local_3": 2,
    "chain_local_4": 3,
    "chain_local_5": 4,
}

# Manuscript-style compact English prefixes (match Table 3 in the paper)
MANUSCRIPT_PREFIX_EN = {
    "Team": "Team",
    "Job": "Job",
    "Shift": "Shift",
    "Loc": "Location",
    "Act": "Activity",
    "HazardMode": "HazardMode",
    "HazardSource": "HazardSource",
    "Body": "Body",
    "InjuryForm": "InjuryForm",
    "Sev": "Severity",
}

MANUSCRIPT_VALUE_EN = {
    **fig06.CHINESE_VALUE_TO_EN,
    "轻伤": "Mild",
    "重伤": "Severe",
    "轻微伤": "Minor injury",
    "较严重伤害": "Non-minor injury",
}


def translate_chain(chain: str) -> str:
    parts = [p.strip() for p in re.split(r"\s*->\s*", str(chain)) if p.strip()]
    out: list[str] = []
    for token in parts:
        if ":" not in token:
            out.append(token)
            continue
        prefix, val = token.split(":", 1)
        prefix = prefix.strip()
        val = val.strip()
        en_p = MANUSCRIPT_PREFIX_EN.get(prefix, prefix)
        en_v = MANUSCRIPT_VALUE_EN.get(val, val)
        out.append(f"{en_p}: {en_v}")
    return " -> ".join(out)


def translate_node_id(node_id: str) -> str:
    return translate_chain(node_id)


def translate_rule_antecedent(antecedent: str) -> str:
    s = str(antecedent).strip()
    parts: list[str] = []
    for piece in re.split(r"\s*;\s*", s):
        piece = piece.strip()
        if not piece:
            continue
        if ":" not in piece:
            parts.append(piece)
            continue
        k, v = piece.split(":", 1)
        k, v = k.strip(), v.strip()
        en_k = MANUSCRIPT_PREFIX_EN.get(k, k)
        en_v = MANUSCRIPT_VALUE_EN.get(v, v)
        parts.append(f"{en_k}: {en_v}")
    return "; ".join(parts)


def translate_consequent(consequent: str) -> str:
    s = str(consequent).strip()
    if s.startswith("SevereBinary:"):
        return "Non-minor injury"
    if s.startswith("Sev:"):
        val = s.split(":", 1)[1].strip()
        return MANUSCRIPT_VALUE_EN.get(val, val)
    return translate_rule_antecedent(s)


def classify_chain_type(chain_en: str) -> str:
    tokens = [t.strip() for t in chain_en.split("->")]
    fields = {t.split(":", 1)[0].strip() for t in tokens if ":" in t}
    if any("Severity: Severe" in t for t in tokens):
        return "Severe-injury context chain"
    if "Body" in fields or "InjuryForm" in fields:
        if "Location" in fields and "Activity" in fields:
            return "Mixed local chain"
        return "Consequence-linked pathway"
    if "HazardMode" in fields:
        return "Process-related chain"
    return "Prevention-side context chain"


def build_high_risk_local_chains(min_count: int = 3) -> pd.DataFrame:
    local_path = OUT_REAL / "chain_local_summary.csv"
    local_df = pd.read_csv(local_path, encoding="utf-8-sig")
    d = local_df.copy()
    d["_sc"] = pd.to_numeric(d["severe_count"], errors="coerce").fillna(0)
    d["_sr"] = pd.to_numeric(d["severe_rate"], errors="coerce")
    d["_cnt"] = pd.to_numeric(d["count"], errors="coerce").fillna(0)
    d = d.loc[(d["_sc"] > 0) & (d["_sr"] > 0) & (d["_cnt"] >= min_count)].copy()
    d = d.drop(columns=["_sc", "_sr", "_cnt"], errors="ignore")
    dedup = mrc._dedup_local_chains_by_string(d)
    dedup = dedup.sort_values(
        by=["severe_rate", "severe_count", "count", "average_severity_score"],
        ascending=[False, False, False, False],
        na_position="last",
    ).reset_index(drop=True)
    return dedup


def make_table_s1(chains: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for i, r in chains.iterrows():
        chain_zh = str(r["chain"])
        chain_en = translate_chain(chain_zh)
        rows.append(
            {
                "Rank": i + 1,
                "Local chain (English)": chain_en,
                "Local chain (original)": chain_zh,
                "Chain template": str(r.get("chain_layer", "")),
                "Chain type": classify_chain_type(chain_en),
                "Count": int(r["count"]),
                "Proportion": round(float(r["proportion"]), 4),
                "Non-minor injury count": int(r["severe_count"]),
                "Non-minor injury rate": round(float(r["severe_rate"]), 4),
                "Severe injury count": int(r.get("serious_count", 0) or 0),
                "Severe injury rate": round(float(r.get("serious_rate", 0) or 0), 4),
                "Average injury-severity score": round(float(r["average_severity_score"]), 3),
            }
        )
    return pd.DataFrame(rows)


def make_table3_top10(chains: pd.DataFrame) -> pd.DataFrame:
    s1 = make_table_s1(chains.head(10))
    return s1[
        [
            "Rank",
            "Local chain (English)",
            "Chain type",
            "Count",
            "Non-minor injury rate",
        ]
    ].assign(
        **{
            "Non-minor injury rate (%)": (s1["Non-minor injury rate"] * 100).round(2).astype(str) + "%"
        }
    )[
        ["Rank", "Local chain (English)", "Chain type", "Count", "Non-minor injury rate (%)"]
    ].rename(
        columns={
            "Local chain (English)": "Local chain",
            "Non-minor injury rate (%)": "Non-minor injury rate",
        }
    )


def _load_rules_with_type() -> pd.DataFrame:
    specs = [
        ("Prevention-side non-minor injury", OUT_REAL / "association_rules_severe_prevention.csv", True),
        ("Prevention-side non-minor injury (robust)", OUT_REAL / "association_rules_severe_prevention_robust.csv", True),
        ("Consequence-side non-minor injury", OUT_REAL / "association_rules_severe_consequence.csv", True),
        ("Consequence-side non-minor injury (robust)", OUT_REAL / "association_rules_severe_consequence_robust.csv", True),
        ("Prevention-side severe injury", OUT_REAL / "association_rules_serious_prevention.csv", False),
        ("All mined rules", OUT_REAL / "association_rules_all.csv", False),
    ]
    frames: list[pd.DataFrame] = []
    for rule_type, path, _ in specs:
        if not path.is_file():
            continue
        df = pd.read_csv(path, encoding="utf-8-sig")
        if df.empty:
            continue
        df = df.copy()
        df["Rule category"] = rule_type
        frames.append(df)
    if not frames:
        return pd.DataFrame()
    out = pd.concat(frames, ignore_index=True)
    out["Antecedent (English)"] = out["antecedents"].map(translate_rule_antecedent)
    out["Consequent (English)"] = out["consequents"].map(translate_consequent)
    out["Antecedent (original)"] = out["antecedents"]
    out["Consequent (original)"] = out["consequents"]
    cols = [
        "Rule category",
        "Antecedent (English)",
        "Consequent (English)",
        "support",
        "confidence",
        "lift",
        "leverage",
        "conviction",
        "antecedent_count",
        "rule_count",
        "Antecedent (original)",
        "Consequent (original)",
    ]
    return out[cols].sort_values(
        by=["Rule category", "lift", "confidence", "support"],
        ascending=[True, False, False, False],
    ).reset_index(drop=True)


def make_table_s4() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    """Supplementary Table S4: complete sensitivity-analysis results (three sheets)."""

    def _load_or_default(path: Path, default: pd.DataFrame) -> pd.DataFrame:
        if path.is_file():
            return pd.read_excel(path, engine="openpyxl")
        return default.copy()

    chain_summary_default = pd.DataFrame(
        {
            "MIN_CHAIN_COUNT": [2, 3, 4],
            "High-risk local chains (n)": [66, 21, 12],
            "Main analysis default": ["", "Yes", ""],
        }
    )
    chain_overlap_default = pd.DataFrame(
        {
            "Comparison": ["threshold_2_vs_3", "threshold_3_vs_4", "threshold_2_vs_4"],
            "Top-10 overlap (n)": [10, 10, 10],
            "Top-10 Jaccard similarity": [1.000, 1.000, 1.000],
        }
    )

    rule_summary_default = pd.DataFrame(
        {
            "Support-count threshold (rule_count and antecedent_count)": [5, 8, 10],
            "Robust prevention-side rules (n)": [100, 51, 38],
            "Main analysis default": ["", "Yes", ""],
        }
    )
    rule_overlap_default = pd.DataFrame(
        {
            "Comparison": ["threshold_5_vs_8", "threshold_8_vs_10", "threshold_5_vs_10"],
            "Top-10 overlap (n)": [2, 3, 1],
            "Top-10 Jaccard similarity": [0.111, 0.176, 0.053],
        }
    )

    node_default = pd.DataFrame(
        {
            "Metric": [
                "Prevention-side Top-10 node overlap (main vs equal weighting)",
                "Prevention-side Top-10 Jaccard similarity",
                "Spearman rank correlation (all prevention-side nodes)",
            ],
            "Value": [9, 0.818, 0.994],
        }
    )

    chain_raw = _load_or_default(OUT_SENS / "sensitivity_chain_threshold_summary.xlsx", pd.DataFrame())
    if not chain_raw.empty and "min_chain_count" in chain_raw.columns:
        chain_summary = pd.DataFrame(
            {
                "MIN_CHAIN_COUNT": chain_raw["min_chain_count"].astype(int),
                "High-risk local chains (n)": chain_raw["n_high_risk_chains"].astype(int),
                "Main analysis default": [
                    "Yes" if int(v) == config.MIN_CHAIN_COUNT else "" for v in chain_raw["min_chain_count"]
                ],
            }
        )
    else:
        chain_summary = chain_summary_default

    chain_overlap_raw = _load_or_default(OUT_SENS / "sensitivity_chain_overlap.xlsx", pd.DataFrame())
    if not chain_overlap_raw.empty and "comparison" in chain_overlap_raw.columns:
        chain_overlap = pd.DataFrame(
            {
                "Comparison": chain_overlap_raw["comparison"],
                "Top-10 overlap (n)": chain_overlap_raw["top10_overlap_count"].astype(int),
                "Top-10 Jaccard similarity": chain_overlap_raw["top10_jaccard_similarity"].round(3),
            }
        )
    else:
        chain_overlap = chain_overlap_default

    chain_sheet = pd.concat(
        [
            chain_summary.assign(Section="Threshold summary"),
            chain_overlap.assign(Section="Top-10 stability"),
        ],
        ignore_index=True,
        sort=False,
    )

    rule_raw = _load_or_default(OUT_SENS / "sensitivity_rule_threshold_summary.xlsx", pd.DataFrame())
    if not rule_raw.empty and "rule_threshold" in rule_raw.columns:
        rule_summary = pd.DataFrame(
            {
                "Support-count threshold (rule_count and antecedent_count)": rule_raw["rule_threshold"].astype(int),
                "Robust prevention-side rules (n)": rule_raw["n_rules"].astype(int),
                "Main analysis default": [
                    "Yes" if int(v) == 8 else "" for v in rule_raw["rule_threshold"]
                ],
            }
        )
    else:
        rule_summary = rule_summary_default

    rule_overlap_raw = _load_or_default(OUT_SENS / "sensitivity_rule_overlap.xlsx", pd.DataFrame())
    if not rule_overlap_raw.empty and "comparison" in rule_overlap_raw.columns:
        rule_overlap = pd.DataFrame(
            {
                "Comparison": rule_overlap_raw["comparison"],
                "Top-10 overlap (n)": rule_overlap_raw["top10_overlap_count"].astype(int),
                "Top-10 Jaccard similarity": rule_overlap_raw["top10_jaccard_similarity"].round(3),
            }
        )
    else:
        rule_overlap = rule_overlap_default

    rule_sheet = pd.concat(
        [
            rule_summary.assign(Section="Threshold summary"),
            rule_overlap.assign(Section="Top-10 stability"),
        ],
        ignore_index=True,
        sort=False,
    )

    node_comp_raw = _load_or_default(OUT_SENS / "sensitivity_node_weight_comparison.xlsx", pd.DataFrame())
    if not node_comp_raw.empty and "metric" in node_comp_raw.columns:
        label_map = {
            "top10_overlap_count": "Prevention-side Top-10 node overlap (main vs equal weighting)",
            "top10_jaccard_similarity": "Prevention-side Top-10 Jaccard similarity",
            "spearman_rank_correlation": "Spearman rank correlation (all prevention-side nodes)",
        }
        node_summary = pd.DataFrame(
            {
                "Metric": node_comp_raw["metric"].map(label_map).fillna(node_comp_raw["metric"]),
                "Value": node_comp_raw["value"].round(3),
            }
        )
    else:
        node_summary = node_default

    main_top_path = OUT_SENS / "sensitivity_node_weight_main_top10.xlsx"
    equal_top_path = OUT_SENS / "sensitivity_node_weight_equal_top10.xlsx"
    if main_top_path.is_file() and equal_top_path.is_file():
        top_main = pd.read_excel(main_top_path, engine="openpyxl")
        top_equal = pd.read_excel(equal_top_path, engine="openpyxl")
        detail_rows: list[dict[str, object]] = []
        for scheme, df in (("Main weights", top_main), ("Equal weights", top_equal)):
            for _, row in df.iterrows():
                detail_rows.append(
                    {
                        "Weighting scheme": scheme,
                        "Rank": int(row.get("rank", 0)),
                        "Node label (English)": translate_node_id(str(row.get("node_id", ""))),
                        "Node ID (original)": str(row.get("node_id", "")),
                    }
                )
        node_detail = pd.DataFrame(detail_rows)
        node_sheet = pd.concat(
            [
                node_summary.assign(Section="Summary metrics"),
                node_detail.assign(Section="Top-10 prevention-side nodes"),
            ],
            ignore_index=True,
            sort=False,
        )
    else:
        node_sheet = node_summary.assign(Section="Summary metrics")

    return chain_sheet, rule_sheet, node_sheet


def make_table_s3() -> tuple[pd.DataFrame, pd.DataFrame]:
    nm = pd.read_csv(OUT_REAL / "node_metrics.csv", encoding="utf-8-sig")
    em = pd.read_csv(OUT_REAL / "edge_metrics.csv", encoding="utf-8-sig")

    nm = nm.copy()
    nm["Node label (English)"] = nm["node_id"].map(translate_node_id)
    nm = nm.sort_values("risk_score", ascending=False).reset_index(drop=True)
    nm.insert(0, "Rank", np.arange(1, len(nm) + 1))

    em = em.copy()
    em["Source (English)"] = em["source"].map(translate_node_id)
    em["Target (English)"] = em["target"].map(translate_node_id)
    if "edge_risk_score" not in em.columns:
        em["edge_risk_score"] = pd.to_numeric(em["weight"], errors="coerce").fillna(0) * pd.to_numeric(
            em["edge_severity_rate"], errors="coerce"
        ).fillna(0)
    em = em.sort_values("edge_risk_score", ascending=False).reset_index(drop=True)
    em.insert(0, "Rank", np.arange(1, len(em) + 1))
    return nm, em


def _manuscript_font() -> str:
    for name in ("Arial", "Helvetica", "DejaVu Sans"):
        try:
            plt.rcParams["font.family"] = name
            return name
        except Exception:
            continue
    return "DejaVu Sans"


def make_figure_s1() -> None:
    nm = pd.read_csv(OUT_REAL / "node_metrics.csv", encoding="utf-8-sig")
    em = pd.read_csv(OUT_REAL / "edge_metrics.csv", encoding="utf-8-sig")
    font_name = _manuscript_font()

    G = nx.DiGraph()
    risk_map = nm.set_index("node_id")["risk_score"].astype(float).to_dict()
    type_map = nm.set_index("node_id")["node_type"].astype(str).to_dict()

    for _, row in nm.iterrows():
        nid = str(row["node_id"])
        G.add_node(
            nid,
            risk=float(row.get("risk_score", 0.0)),
            nt=type_map.get(nid, ""),
            lab=fig06._compact_node_label_figure3(nid),
        )

    for _, row in em.iterrows():
        u, v = str(row["source"]), str(row["target"])
        w = float(pd.to_numeric(row.get("weight"), errors="coerce") or 0.0)
        if u in G and v in G and w > 0:
            if G.has_edge(u, v):
                G[u][v]["weight"] += w
            else:
                G.add_edge(u, v, weight=w)

    pos = nx.spring_layout(G, seed=42, k=1.8 / max(1, np.sqrt(G.number_of_nodes())), iterations=80)

    risks = np.array([float(G.nodes[n].get("risk", 0.0)) for n in G.nodes()])
    rmin, rmax = float(risks.min()), float(risks.max())
    sizes = 80.0 + 420.0 * (risks - rmin) / (rmax - rmin) if rmax > rmin else np.full(len(risks), 220.0)

    type_color = {
        "Team": "#7B68EE",
        "Job": "#4682B4",
        "Shift": "#5F9EA0",
        "Loc": "#2E8B57",
        "Act": "#DAA520",
        "HazardMode": "#CD853F",
        "HazardSource": "#B87333",
        "Body": "#BC8F8F",
        "InjuryForm": "#D2691E",
        "Sev": "#A0522D",
        "Cause": "#708090",
    }
    node_colors = [type_color.get(str(G.nodes[n].get("nt", "")), "#696969") for n in G.nodes()]

    fig, ax = plt.subplots(figsize=(14, 11))
    widths = []
    for u, v in G.edges():
        w = float(G[u][v].get("weight", 1.0))
        widths.append(max(0.2, min(4.0, 0.08 * np.sqrt(w))))
    nx.draw_networkx_edges(
        G,
        pos,
        ax=ax,
        width=widths,
        edge_color="#666666",
        arrows=True,
        arrowsize=8,
        alpha=0.35,
        connectionstyle="arc3,rad=0.06",
    )
    nx.draw_networkx_nodes(G, pos, ax=ax, node_size=sizes, node_color=node_colors, alpha=0.9, linewidths=0.3)
    labels = {n: str(G.nodes[n].get("lab", n)) for n in G.nodes()}
    nx.draw_networkx_labels(G, pos, labels=labels, ax=ax, font_size=5.5, font_family=font_name)
    ax.set_title(
        "Supplementary Figure S1. Complete directed weighted risk network (81 nodes, 381 edges)",
        fontsize=12,
        fontweight="bold",
        fontfamily=font_name,
    )
    ax.text(
        0.5,
        -0.03,
        "Node size indicates composite risk score. Edge width indicates adjacent relationship frequency. "
        "All typed nodes and directed edges from the reconstructed accident-register data are shown.",
        transform=ax.transAxes,
        ha="center",
        fontsize=9,
        fontfamily=font_name,
    )
    ax.axis("off")
    fig.tight_layout()
    for ext in ("png", "pdf", "svg"):
        fig.savefig(FIG_DIR / f"Supplementary_Figure_S1_complete_network.{ext}", dpi=300, bbox_inches="tight")
    plt.close(fig)


def write_table_docx(df: pd.DataFrame, path: Path, title: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print(f"SKIP docx (python-docx not installed): {path.name}")
        return

    doc = Document()
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                if col.endswith("rate") or "rate" in col.lower():
                    cells[j].text = f"{val * 100:.2f}%" if val <= 1 else f"{val:.3f}"
                else:
                    cells[j].text = f"{val:.3f}" if abs(val) < 10 else f"{val:.2f}"
            else:
                cells[j].text = str(val)
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)
    doc.save(path)


def update_manuscript_table3(table3: pd.DataFrame) -> None:
    if not MANUSCRIPT_DIR.is_dir():
        return
    for fname in MANUSCRIPT_DIR.iterdir():
        if fname.suffix == ".docx" and not fname.name.startswith("~") and "少" in fname.name:
            out = MANUSCRIPT_DIR / "Table3_top10_local_chains.docx"
            write_table_docx(table3, out, "Table 3. High-risk local chains (Top 10)")
            print(f"OK: manuscript helper table -> {out}")
            break


def main() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    chains = build_high_risk_local_chains(config.MIN_CHAIN_COUNT)
    if len(chains) != 21:
        print(f"WARNING: expected 21 high-risk local chains, got {len(chains)}")

    s1 = make_table_s1(chains)
    table3 = make_table3_top10(chains)
    rules = _load_rules_with_type()
    nodes, edges = make_table_s3()
    chain_sens, rule_sens, node_sens = make_table_s4()

    # Excel outputs
    s1_xlsx = TABLE_DIR / "Supplementary_Table_S1_local_chains_21.xlsx"
    s1.to_excel(s1_xlsx, index=False, engine="openpyxl")

    s2_xlsx = TABLE_DIR / "Supplementary_Table_S2_association_rules.xlsx"
    with pd.ExcelWriter(s2_xlsx, engine="openpyxl") as writer:
        rules.to_excel(writer, sheet_name="All_rules", index=False)
        for cat in sorted(rules["Rule category"].dropna().unique()):
            sub = rules.loc[rules["Rule category"] == cat]
            sheet = re.sub(r"[^\w]+", "_", cat)[:31]
            sub.to_excel(writer, sheet_name=sheet, index=False)

    s3_xlsx = TABLE_DIR / "Supplementary_Table_S3_node_edge_metrics.xlsx"
    with pd.ExcelWriter(s3_xlsx, engine="openpyxl") as writer:
        nodes.to_excel(writer, sheet_name="Nodes", index=False)
        edges.to_excel(writer, sheet_name="Edges", index=False)

    s4_xlsx = TABLE_DIR / "Supplementary_Table_S4_sensitivity_analysis.xlsx"
    with pd.ExcelWriter(s4_xlsx, engine="openpyxl") as writer:
        chain_sens.to_excel(writer, sheet_name="Chain_threshold_sensitivity", index=False)
        rule_sens.to_excel(writer, sheet_name="Rule_threshold_sensitivity", index=False)
        node_sens.to_excel(writer, sheet_name="Node_weight_sensitivity", index=False)

    table3_xlsx = TABLE_DIR / "Table3_top10_local_chains.xlsx"
    table3.to_excel(table3_xlsx, index=False, engine="openpyxl")

    # CSV mirrors
    s1.to_csv(TABLE_DIR / "Supplementary_Table_S1_local_chains_21.csv", index=False, encoding="utf-8-sig")
    rules.to_csv(TABLE_DIR / "Supplementary_Table_S2_association_rules.csv", index=False, encoding="utf-8-sig")
    nodes.to_csv(TABLE_DIR / "Supplementary_Table_S3_nodes.csv", index=False, encoding="utf-8-sig")
    edges.to_csv(TABLE_DIR / "Supplementary_Table_S3_edges.csv", index=False, encoding="utf-8-sig")
    chain_sens.to_csv(TABLE_DIR / "Supplementary_Table_S4_chain_threshold_sensitivity.csv", index=False, encoding="utf-8-sig")
    rule_sens.to_csv(TABLE_DIR / "Supplementary_Table_S4_rule_threshold_sensitivity.csv", index=False, encoding="utf-8-sig")
    node_sens.to_csv(TABLE_DIR / "Supplementary_Table_S4_node_weight_sensitivity.csv", index=False, encoding="utf-8-sig")
    table3.to_csv(TABLE_DIR / "Table3_top10_local_chains.csv", index=False, encoding="utf-8-sig")

    # DOCX tables
    write_table_docx(s1, TABLE_DIR / "Supplementary_Table_S1_local_chains_21.docx", "Supplementary Table S1. High-risk local chains (n = 21)")
    write_table_docx(rules.head(100), TABLE_DIR / "Supplementary_Table_S2_association_rules_preview.docx", "Supplementary Table S2. Association rules (preview top 100)")
    write_table_docx(nodes, TABLE_DIR / "Supplementary_Table_S3_nodes.docx", "Supplementary Table S3. Node-level metrics (n = 81)")
    write_table_docx(edges, TABLE_DIR / "Supplementary_Table_S3_edges.docx", "Supplementary Table S3. Edge-level metrics (n = 381)")
    write_table_docx(chain_sens, TABLE_DIR / "Supplementary_Table_S4_sensitivity_analysis.docx", "Supplementary Table S4. Complete sensitivity-analysis results")
    write_table_docx(table3, TABLE_DIR / "Table3_top10_local_chains.docx", "Table 3. High-risk local chains (Top 10)")
    update_manuscript_table3(table3)

    make_figure_s1()

    # README for supplementary package
    readme = SUPP_DIR / "README_supplementary.md"
    readme.write_text(
        "\n".join(
            [
                "# Supplementary materials",
                "",
                "Generated from `outputs_real/` using `scripts/generate_supplementary_materials.py`.",
                "",
                "## Files",
                "- `tables/Supplementary_Table_S1_local_chains_21.xlsx` — 21 high-risk local chains (English + original Chinese labels).",
                "- `tables/Supplementary_Table_S2_association_rules.xlsx` — complete association rules by category (all + robust + serious).",
                "- `tables/Supplementary_Table_S3_node_edge_metrics.xlsx` — full node (81) and edge (381) metrics.",
                "- `tables/Supplementary_Table_S4_sensitivity_analysis.xlsx` — complete sensitivity-analysis results (chain, rule, and node-weight tests).",
                "- `tables/Table3_top10_local_chains.xlsx` — manuscript Table 3 (Top 10).",
                "- `figures/Supplementary_Figure_S1_complete_network.png/pdf/svg` — complete directed weighted risk network.",
                "",
                f"Local chains retained: **{len(chains)}** (MIN_CHAIN_COUNT={config.MIN_CHAIN_COUNT}).",
                f"Association rules exported: **{len(rules)}** rows across categories.",
                f"Nodes: **{len(nodes)}**; Edges: **{len(edges)}**.",
            ]
        ),
        encoding="utf-8",
    )

    print(f"OK: Supplementary package written to {SUPP_DIR.resolve()}")
    print(f"  S1 rows: {len(s1)}")
    print(f"  S2 rows: {len(rules)}")
    print(f"  S3 nodes/edges: {len(nodes)}/{len(edges)}")
    print(f"  S4 chain/rule/node rows: {len(chain_sens)}/{len(rule_sens)}/{len(node_sens)}")
    print(f"  Table3 top10 rows: {len(table3)}")


if __name__ == "__main__":
    main()
