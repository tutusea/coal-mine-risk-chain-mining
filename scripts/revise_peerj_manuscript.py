# -*- coding: utf-8 -*-
"""Create a PeerJ Computer Science-oriented revision of manus.docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent.parent
MANUSCRIPT_DIR = ROOT / "1正文中文" / "1Chiang Mai Journal of Science"
SRC = MANUSCRIPT_DIR / "manus.docx"
OUT = MANUSCRIPT_DIR / "manus_PeerJ_CS_revised.docx"


def set_paragraph_text(paragraph, text: str):
    style = paragraph.style
    alignment = paragraph.alignment
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    paragraph.style = style
    paragraph.alignment = alignment
    return run


def bold_abstract_labels(paragraph) -> None:
    text = paragraph.text
    labels = ["Background.", "Methods.", "Results.", "Conclusions."]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    remainder = text
    segments: list[tuple[bool, str]] = []
    for label in labels:
        pos = remainder.find(label)
        if pos > 0:
            segments.append((False, remainder[:pos]))
            remainder = remainder[pos:]
        if remainder.startswith(label):
            segments.append((True, label))
            remainder = remainder[len(label) :]
    if remainder:
        segments.append((False, remainder))

    for is_label, segment in segments:
        run = paragraph.add_run(segment)
        run.bold = is_label


def main() -> None:
    doc = Document(str(SRC))

    set_paragraph_text(
        doc.paragraphs[0],
        "Interpretable Semantic Reconstruction and Graph-Based Local Risk-Chain Mining "
        "for Small-Sample Occupational Injury Records: A Coal Mine Case Study",
    )

    set_paragraph_text(
        doc.paragraphs[6],
        "Background. Small occupational injury registers are difficult to analyze with conventional "
        "prediction-oriented models because their fields are often semantically mixed, sparse, and "
        "closely coupled with injury outcomes. Methods. This study developed an interpretable "
        "graph-based mining framework for small-sample accident-register data. A total of 234 real "
        "coal mine injury records were semantically reconstructed by separating the original "
        "hazard-related field into hazard mode, injury form, and hazard source. The reconstructed "
        "records were encoded as typed risk nodes, transformed into local risk chains, and represented "
        "as a directed weighted risk network. Node, edge, chain, and association-rule analyses were "
        "then combined with sensitivity analysis to evaluate the stability of the extracted risk "
        "evidence. Results. The reconstructed network contained 81 nodes and 381 directed edges. "
        "After applying frequency, non-minor injury, and deduplication criteria, 21 high-risk local "
        "chains were retained. Transportation, coalface work, middle shift, tunneling workers, and "
        "coal miners repeatedly appeared in prevention-side nodes, local chains, or robust association "
        "rules associated with non-minor injuries. Sensitivity analysis showed that the top local "
        "chains and prevention-side node rankings were stable across chain-frequency thresholds and "
        "weighting schemes, whereas association-rule rankings were more sensitive to support-count "
        "thresholds. Conclusions. The proposed framework provides a reproducible and interpretable "
        "approach for extracting structured risk evidence from small, heterogeneous occupational "
        "injury registers. The coal mine case study illustrates how semantic reconstruction and "
        "graph-based local risk-chain mining can reduce outcome leakage, preserve pathway information, "
        "and support cautious interpretation without claiming causal inference.",
    )
    bold_abstract_labels(doc.paragraphs[6])

    set_paragraph_text(
        doc.paragraphs[7],
        "Keywords: interpretable data mining; semantic reconstruction; graph-based local risk-chain "
        "mining; directed weighted risk network; association-rule mining; occupational injury records; "
        "small-sample data",
    )

    set_paragraph_text(
        doc.paragraphs[14],
        "To address these problems, this study proposes an interpretable graph-based data-mining "
        "framework for small-sample occupational injury records. The framework first reconstructs the "
        "semantically mixed hazard-related field into three analytical variables: HazardMode, "
        "InjuryForm, and HazardSource. This step separates injury-process information from "
        "injury-outcome information and reduces the risk of treating consequences as pre-accident risk "
        "factors. The reconstructed records are then encoded as typed risk nodes and transformed into "
        "local risk chains. A directed weighted risk network is constructed from adjacent node "
        "relationships, and node, edge, chain, and association-rule analyses are used to extract "
        "complementary forms of interpretable risk evidence. Sensitivity analysis is further used to "
        "evaluate whether chain, rule, and node-ranking results are stable under alternative parameter "
        "settings.",
    )

    set_paragraph_text(
        doc.paragraphs[15],
        "The objectives of this study are: (1) to formalize a semantic reconstruction strategy for "
        "semantically mixed accident-register fields; (2) to construct local risk chains that reduce "
        "sparsity in small accident datasets while preserving ordered pathway information; (3) to build "
        "a directed weighted risk network and identify prevention-side nodes and edge relationships; "
        "(4) to mine robust association rules related to non-minor injuries while controlling the "
        "interpretation boundary between prevention-side and consequence-side variables; and (5) to "
        "assess the stability of the main findings under different chain-frequency thresholds, rule "
        "support-count thresholds, and node-weighting schemes.",
    )

    set_paragraph_text(
        doc.paragraphs[16],
        "The main contribution of this study is methodological rather than purely domain-specific. It "
        "provides an interpretable and reproducible graph-mining workflow for transforming small, "
        "heterogeneous accident-register records into structured local chains, network evidence, and "
        "threshold-tested co-occurrence patterns. Although coal mine injury records are used as the "
        "empirical case, the framework is intended for occupational safety datasets in which sample "
        "sizes are limited, fields are semantically mixed, and explainable risk interpretation is more "
        "important than black-box prediction accuracy.",
    )

    doc.paragraphs[21].insert_paragraph_before(
        "Data governance and ethics. The analysis used retrospective, de-identified administrative "
        "injury records. The records available for analysis did not include workers' names, "
        "identification numbers, contact information, employee numbers, mine identity, exact accident "
        "dates, or other directly identifiable personal information. Because the source records contain "
        "sensitive organization-specific occupational injury information, the raw administrative "
        "register cannot be publicly released. The reproducible materials for peer review should "
        "therefore include the de-identified derived analytical tables, field dictionary, configuration "
        "file, analysis scripts, and output tables used to generate the figures and supplementary "
        "materials."
    )

    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "2.4 Directed weighted risk network":
            doc.paragraphs[idx].insert_paragraph_before(
                "Computational reproducibility. The analytical workflow was implemented as a "
                "deterministic Python pipeline. The main steps were data cleaning and real-sample "
                "selection, semantic reconstruction, local risk-chain construction, directed weighted "
                "risk-network construction, association-rule mining, sensitivity analysis, and "
                "figure/table generation. All thresholds used in the main analysis and sensitivity "
                "analysis were stored as explicit configuration values so that the retained chains, "
                "rules, node rankings, figures, and supplementary tables could be regenerated from the "
                "same derived input tables."
            )
            break

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "3. Results and Discussion":
            set_paragraph_text(paragraph, "3. Results")
        elif text == "3.7 Representative pattern interpretation and methodological implications":
            set_paragraph_text(paragraph, "4. Discussion")
        elif text == "3.8 Limitations and future work":
            set_paragraph_text(paragraph, "4.2 Limitations and future work")
        elif text == "4. Conclusions":
            set_paragraph_text(paragraph, "5. Conclusions")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(
            "To further illustrate how the proposed framework supports interpretation"
        ):
            set_paragraph_text(
                paragraph,
                "The results should be interpreted primarily as evidence for an explainable "
                "graph-mining workflow rather than as a universal description of coal mine injury "
                "mechanisms. To illustrate how the proposed framework supports interpretation, "
                "representative mined patterns were examined together with their corresponding "
                "work-context and injury-outcome structures. The purpose of this step was not to "
                "conduct a separate qualitative case study, but to evaluate whether the mined local "
                "chains could be traced back to meaningful accident scenarios and whether their "
                "prevention-side and consequence-side meanings could be separated.",
            )
            paragraph.insert_paragraph_before(
                "4.1 Methodological implications and interpretation boundaries"
            )
            break

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(
            "The raw administrative accident-register records contain organization-specific"
        ):
            set_paragraph_text(
                paragraph,
                "The raw administrative accident-register records contain organization-specific and "
                "potentially sensitive occupational injury information and cannot be made publicly "
                "available because public release could compromise institutional confidentiality and "
                "data-use restrictions. De-identified derived analytical tables, field dictionaries, "
                "configuration files, and output tables supporting the figures and supplementary "
                "materials should be deposited in an immutable public repository before submission "
                "(for example, Zenodo or Figshare; DOI: [to be added before submission]) or provided "
                "as PeerJ Supplemental Files where repository release is not permitted. These derived "
                "materials should be sufficient to reproduce the reported node, edge, chain, "
                "association-rule, and sensitivity-analysis results without disclosing the original "
                "administrative register.",
            )
        elif text.startswith("The Python scripts used for semantic reconstruction"):
            set_paragraph_text(
                paragraph,
                "The Python scripts used for data cleaning, semantic reconstruction, risk-chain "
                "construction, network analysis, association-rule mining, sensitivity analysis, and "
                "figure/table generation should be made available in a version-controlled repository "
                "archived with a persistent DOI before submission (DOI: [to be added before "
                "submission]). The repository should include the configuration file, "
                "environment/dependency information, and instructions for regenerating the main "
                "outputs and supplementary tables from the de-identified derived analytical tables.",
            )
        elif text.startswith("During the preparation of this manuscript, generative AI"):
            set_paragraph_text(
                paragraph,
                "During the preparation of this manuscript, generative AI and AI-assisted technologies "
                "were used only to improve language clarity, readability, and editorial presentation. "
                "The tools were not used to generate or alter the underlying data, statistical results, "
                "figures, or scientific conclusions. The authors reviewed, revised, and verified all "
                "AI-assisted outputs and take full responsibility for the accuracy and integrity of the "
                "manuscript, including references, figures, and supplementary materials.",
            )
        elif text.startswith("The supplementary material includes Supplementary Table S1"):
            set_paragraph_text(
                paragraph,
                "The supplementary material includes Supplementary Table S1, the complete list of "
                "retained high-risk local chains; Supplementary Table S2, complete association-rule "
                "outputs for prevention-side non-minor injury rules, consequence-side non-minor injury "
                "rules, and prevention-side severe-injury rules; Supplementary Table S3, complete node "
                "and edge metrics of the directed weighted risk network; Supplementary Table S4, "
                "complete sensitivity-analysis results under different chain-frequency thresholds, "
                "support-count thresholds, and node-weighting schemes; and Supplementary Figure S1, "
                "the complete directed weighted risk-network visualization. For PeerJ Computer Science "
                "submission, the supplementary package should also be accompanied by the de-identified "
                "derived analytical tables and code repository described in the Data availability and "
                "Code availability statements.",
            )

    for run in doc.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(14)

    doc.save(str(OUT))
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
